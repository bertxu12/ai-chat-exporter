import streamlit as st
import re
from datetime import datetime
from io import BytesIO

# Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Excel
import pandas as pd


def clean_special_chars(text: str, aggressive: bool = False, preserve_code: bool = True) -> str:
    """
    修复版清理函数 - 避免删除代码块，保留JS/HTML等内容
    """
    if not text:
        return text

    # 先提取并保护多行代码块（用特殊标记包裹，防止后续正则干扰）
    code_blocks = []

    def replace_code_block(match):
        code_content = match.group(0)[3:-3].strip()  # 去掉```和语言标识，保留纯内容
        code_blocks.append(code_content)
        return f"< preserved_code_{len(code_blocks) - 1} >"  # 临时占位符

    if preserve_code:
        text = re.sub(r'```[\s\S]*?```', replace_code_block, text)
    else:
        # 如果不保留，直接删除（原有行为）
        text = re.sub(r'```[\s\S]*?```', '', text)

    # 2. 行内代码 → 只保留内容（但保留在上下文中）
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # 3. 链接 → 只保留显示文字
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'!\[([^\]]*)\]\([^)]*\)', r'\1', text)

    # 4. 标题符号
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    # 5. 列表符号 → 转成缩进（不直接删除内容）
    text = re.sub(r'^\s*([-*+•◦➤]|(\d+[.)]))\s+', '  • ', text, flags=re.MULTILINE)

    # 6. 清理强调、删除线 - 更安全版本
    for mark in [r'\*{1,3}', r'_{1,2}', r'~~']:
        pattern = rf'({mark})(.+?)({mark})(?!\S)'
        text = re.sub(pattern, r'\2', text, flags=re.DOTALL)

    # 清理孤立标记
    text = re.sub(r'\*{2,3}|_{2,3}|~~|\*\*', '', text)

    # 7. 移除表情符号和常见装饰字符
    text = re.sub(
        r'[\U0001F300-\U0001F9FF\U0001FA00-\U0001FAFF'
        r'\U00002700-\U000027BF\U00002600-\U000026FF'
        r'\U0001F000-\U0001FFFF]+', '', text)
    text = re.sub(r'[★☆♡♥♦♠♣●○◆◇■□▲△▼▽◀▶※♪♫✓✔✕✖]', '', text)

    # 8. 激进模式（只做最必要的过滤）
    if aggressive:
        text = re.sub(
            r'[^\u4e00-\u9fffa-zA-Z0-9\s'
            r'\u3000-\u303F\uFF00-\uFFEF'  # 中文标点 + 全角
            r'。，、；：？！…—～·（）【】《》""''\'\"-.,;:!?()%+*/=&@#$^]',
            '', text)

    # 9. 收尾规范化
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # 压缩多空行
    text = re.sub(r'[ \t]{2,}', ' ', text)  # 多空格 → 单空格
    text = re.sub(r'\s+([，。、；：？！）】》"])', r'\1', text)  # 中文标点前去空格

    # 最后，放回保护的代码块（可选：添加换行和缩进以保持可读性）
    for i, code in enumerate(code_blocks):
        formatted_code = '\n'.join('    ' + line for line in code.split('\n'))  # 添加缩进，模拟代码格式
        text = text.replace(f"< preserved_code_{i} >", f"\n[代码块]\n{formatted_code}\n[/代码块]")

    return text.strip()


def parse_dialog(text: str) -> list:
    """对话解析 - 保持不变"""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    messages = []
    current_role = None
    current_content = []

    user_keywords = {'用户', '我', 'user', 'me', 'human'}
    ai_keywords = {'ai', 'grok', 'claude', 'chatgpt', 'gpt', 'assistant', '助手', 'bot'}

    for line in lines:
        role = None
        content = line

        if '：' in line or ':' in line:
            sep = '：' if '：' in line else ':'
            parts = line.split(sep, 1)
            role_part = parts[0].strip().lower()
            content_part = parts[1].strip() if len(parts) > 1 else ''

            if any(k in role_part for k in user_keywords):
                role = 'user'
                content = content_part
            elif any(k in role_part for k in ai_keywords):
                role = 'assistant'
                content = content_part

        if role:
            if current_role and current_content:
                messages.append({
                    'role': current_role,
                    'content': '\n'.join(current_content).strip()
                })
            current_role = role
            current_content = [content] if content else []
        else:
            if current_role:
                current_content.append(line)
            else:
                current_role = 'user'
                current_content = [line]

    if current_role and current_content:
        messages.append({
            'role': current_role,
            'content': '\n'.join(current_content).strip()
        })

    return messages


def parse_markdown_tables(content):
    """
    解析 content 中的所有 Markdown 表格，返回列表：每个元素是 (pre_text, header, rows, post_text)
    如果没有表格，返回 [(content, None, None, '')]
    """
    parts = []
    last_end = 0

    # 匹配 Markdown 表格：表头 | 分隔符 | 数据行
    # 允许表格前后有其他内容
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # 检测表头（以 | 开头和结尾）
        if line and line.startswith('|') and line.endswith('|'):
            # 可能是表格的开始
            table_start = i
            header_line = line

            # 检查下一行是否是分隔符
            if i + 1 < len(lines):
                separator_line = lines[i + 1].strip()
                # 分隔符行应该包含 - 和 |
                if separator_line and '|' in separator_line and ('-' in separator_line or ':' in separator_line):
                    # 这是一个表格！
                    # 解析表头
                    header = [cell.strip() for cell in header_line.split('|')[1:-1]]

                    # 收集数据行
                    data_rows = []
                    j = i + 2
                    while j < len(lines):
                        row_line = lines[j].strip()
                        # 检查是否是表格行
                        if row_line and row_line.startswith('|') and row_line.endswith('|'):
                            row_cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                            # 确保列数匹配
                            if len(row_cells) == len(header):
                                data_rows.append(row_cells)
                                j += 1
                            else:
                                break
                        else:
                            break

                    # 提取前置文本
                    pre_text = '\n'.join(lines[last_end:table_start]).strip()

                    # 记录表格
                    parts.append((pre_text, header, data_rows, ''))

                    # 更新位置
                    last_end = j
                    i = j
                    continue

        i += 1

    # 处理最后的后置文本
    post_text = '\n'.join(lines[last_end:]).strip()

    if parts:
        # 将后置文本添加到最后一个 part
        if post_text:
            pre, h, r, _ = parts[-1]
            parts[-1] = (pre, h, r, post_text)
    else:
        # 没有找到表格
        parts.append((content, None, None, ''))

    return parts


def generate_word(messages, title):
    doc = Document()

    # 标题
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(22)
    run.bold = True

    # 元信息
    p = doc.add_paragraph(
        f"导出时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}\n"
        f"消息数量：{len(messages)} 条"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    for i, msg in enumerate(messages, 1):
        p = doc.add_paragraph()
        role_text = f"用户（第 {i} 轮）" if msg['role'] == 'user' else f"AI助手（第 {i} 轮）"
        run = p.add_run(role_text)
        run.bold = True
        run.font.size = Pt(14)

        if msg['role'] == 'user':
            run.font.color.rgb = RGBColor(37, 99, 235)
        else:
            run.font.color.rgb = RGBColor(22, 163, 74)

        # 解析表格和代码
        parts = parse_markdown_tables(msg['content'])
        for pre_text, header, rows, post_text in parts:
            if pre_text:
                content_parts = re.split(r'\[代码块\](.*?)\[/代码块\]', pre_text, flags=re.DOTALL)
                for part in content_parts:
                    if part.strip():
                        p = doc.add_paragraph(part)
                        if re.match(r'^\s*\n', part):  # 代码
                            for run in p.runs:
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                            p.paragraph_format.left_indent = Pt(20)  # 缩进

            if header and rows:
                table = doc.add_table(rows=len(rows) + 1, cols=len(header))
                table.style = 'Table Grid'  # 使用网格样式
                hdr_cells = table.rows[0].cells
                for j, h in enumerate(header):
                    hdr_cells[j].text = h
                    hdr_cells[j].paragraphs[0].runs[0].bold = True
                    hdr_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for row_idx, row_data in enumerate(rows, 1):
                    row_cells = table.rows[row_idx].cells
                    for j, cell_text in enumerate(row_data):
                        row_cells[j].text = cell_text
                        row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 调整列宽（可选）
                for column in table.columns:
                    column.width = Inches(2.0)  # 根据需要调整

            if post_text:
                content_parts = re.split(r'\[代码块\](.*?)\[/代码块\]', post_text, flags=re.DOTALL)
                for part in content_parts:
                    if part.strip():
                        p = doc.add_paragraph(part)
                        if re.match(r'^\s*\n', part):  # 代码
                            for run in p.runs:
                                run.font.name = 'Courier New'
                                run.font.size = Pt(10)
                            p.paragraph_format.left_indent = Pt(20)  # 缩进

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_excel(messages, title, pure_mode=False):
    """
    导出Excel，支持两种模式：
    1. 完整模式（pure_mode=False）：包含轮次、角色等信息
    2. 纯表格模式（pure_mode=True）：只导出表格数据，去除所有元数据
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    buffer = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "对话记录"

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    if pure_mode:
        # 纯表格模式：只保留表格数据
        current_row = 1
        first_table = True

        for msg in messages:
            parts = parse_markdown_tables(msg['content'])

            for pre_text, header, rows, post_text in parts:
                if header and rows:
                    # 如果不是第一个表格，空两行
                    if not first_table:
                        current_row += 2
                    first_table = False

                    # 写入表头
                    for col_idx, col_name in enumerate(header, 1):
                        cell = ws.cell(row=current_row, column=col_idx)
                        cell.value = col_name
                        cell.font = Font(bold=True, size=11)
                        cell.fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                        cell.border = thin_border
                        ws.column_dimensions[cell.column_letter].width = 12

                    current_row += 1

                    # 写入数据行
                    for row_data in rows:
                        for col_idx, cell_value in enumerate(row_data, 1):
                            cell = ws.cell(row=current_row, column=col_idx)
                            cell.value = cell_value
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                            cell.border = thin_border
                        current_row += 1

        wb.save(buffer)

    else:
        # 完整模式：包含对话信息
        current_row = 1

        # 设置表头
        headers = ['轮次', '角色', '内容', '字数']
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.value = header
            cell.font = Font(bold=True, size=12)
            cell.fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # 设置列宽
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 100
        ws.column_dimensions['D'].width = 10

        current_row = 2

        for i, msg in enumerate(messages, 1):
            # 解析内容
            content = msg['content']
            parts = parse_markdown_tables(content)

            # 组装显示内容
            display_parts = []

            for pre_text, header, rows, post_text in parts:
                if pre_text:
                    display_parts.append(pre_text)

                if header and rows:
                    # 将表格转换为易读的文本格式
                    table_text = ' | '.join(header) + '\n'
                    table_text += '-' * 50 + '\n'
                    for row in rows:
                        table_text += ' | '.join(row) + '\n'
                    display_parts.append(table_text.strip())

                if post_text:
                    display_parts.append(post_text)

            clean_content = '\n\n'.join(display_parts)

            # 写入数据
            ws.cell(row=current_row, column=1).value = i
            ws.cell(row=current_row, column=1).alignment = Alignment(horizontal='center', vertical='center')

            ws.cell(row=current_row, column=2).value = '用户' if msg['role'] == 'user' else 'AI助手'
            ws.cell(row=current_row, column=2).alignment = Alignment(horizontal='center', vertical='center')

            content_cell = ws.cell(row=current_row, column=3)
            content_cell.value = clean_content
            content_cell.alignment = Alignment(wrap_text=True, vertical='top', horizontal='left')

            ws.cell(row=current_row, column=4).value = len(clean_content)
            ws.cell(row=current_row, column=4).alignment = Alignment(horizontal='center', vertical='center')

            current_row += 1

        # 添加边框
        for row in ws.iter_rows(min_row=1, max_row=current_row - 1, min_col=1, max_col=4):
            for cell in row:
                cell.border = thin_border

        wb.save(buffer)

    buffer.seek(0)
    return buffer


def main():
    st.set_page_config(page_title="AI对话导出工具", page_icon="💬", layout="wide")

    if "original_text" not in st.session_state:
        st.session_state.original_text = ""
    if "current_text" not in st.session_state:
        st.session_state.current_text = ""
    if "cleaned_once" not in st.session_state:
        st.session_state.cleaned_once = False

    st.title("AI 对话导出工具")
    st.caption("支持 Word 和 Excel 格式导出")

    with st.sidebar:
        st.header("导出设置")
        title = st.text_input("文档标题", "AI对话记录")

        st.divider()
        st.subheader("导出格式")
        export_word = st.checkbox("Word (.docx)", True)
        export_excel = st.checkbox("Excel (.xlsx)", False)

        # Excel 表格选项
        if export_excel:
            st.markdown("**Excel 选项：**")
            excel_mode = st.radio(
                "导出模式",
                options=["完整模式（包含轮次/角色）", "纯表格模式（仅保留表格数据）"],
                index=1,
                help="完整模式：包含对话的轮次、角色等信息\n纯表格模式：只导出表格内容，去除所有元数据"
            )
            excel_pure_mode = (excel_mode == "纯表格模式（仅保留表格数据）")
        else:
            excel_pure_mode = False

        st.divider()
        st.subheader("文本清理")
        auto_clean = st.checkbox("导出时自动清理", True)
        aggressive = st.checkbox("激进模式（最大程度去干扰）", False)
        preserve_code = st.checkbox("保留代码块（如JS/HTML）", True)

    col1, col2 = st.columns([2, 1])

    with col1:
        st.subheader("对话内容")

        raw_text = st.text_area(
            "请粘贴完整对话...",
            value=st.session_state.current_text,
            height=500
        )

        st.session_state.current_text = raw_text

        btn_col1, btn_col2, btn_col3 = st.columns([1, 1, 2])

        with btn_col1:
            if st.button("🧹 清理文本", type="primary"):
                if raw_text.strip():
                    if not st.session_state.cleaned_once:
                        st.session_state.original_text = raw_text
                    cleaned = clean_special_chars(raw_text, aggressive=aggressive, preserve_code=preserve_code)
                    st.session_state.current_text = cleaned
                    st.session_state.cleaned_once = True
                    st.success("清理完成")
                    st.rerun()

        with btn_col2:
            if st.button("↩️ 恢复原始"):
                if st.session_state.original_text:
                    st.session_state.current_text = st.session_state.original_text
                    st.session_state.cleaned_once = False
                    st.rerun()

        with btn_col3:
            if st.button("🗑️ 清空"):
                for key in ["current_text", "original_text", "cleaned_once"]:
                    if key in st.session_state:
                        del st.session_state[key]
                st.rerun()

    with col2:
        st.subheader("统计信息")
        if st.session_state.current_text.strip():
            messages = parse_dialog(st.session_state.current_text)
            if messages:
                st.metric("消息数量", len(messages))
                st.metric("总字符数", f"{sum(len(m['content']) for m in messages):,}")

    # 导出部分
    if st.session_state.current_text.strip():
        messages = parse_dialog(st.session_state.current_text)

        final_messages = messages
        if auto_clean:
            final_messages = []
            for m in messages:
                cleaned = clean_special_chars(m['content'], aggressive=aggressive, preserve_code=preserve_code)
                final_messages.append({'role': m['role'], 'content': cleaned})

        if messages and (export_word or export_excel):
            st.divider()
            st.subheader("导出")

            cols = st.columns(2)

            with cols[0]:
                if export_word and st.button("生成 Word"):
                    buf = generate_word(final_messages, title)
                    st.download_button(
                        "⬇️ 下载 Word", buf,
                        f"{title}_{datetime.now():%Y%m%d_%H%M}.docx",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            with cols[1]:
                if export_excel and st.button("生成 Excel"):
                    buf = generate_excel(final_messages, title, pure_mode=excel_pure_mode)
                    st.download_button(
                        "⬇️ 下载 Excel", buf,
                        f"{title}_{datetime.now():%Y%m%d_%H%M}.xlsx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )


if __name__ == "__main__":
    main()